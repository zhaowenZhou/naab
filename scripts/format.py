import docx
import docx2txt
import re
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

def reformat(file_path):
    text = docx2txt.process(file_path)
    text = text.replace('\n', '。').replace(' ', '，')
    # 处理重复的标点符号
    text = re.sub(r'([，。；！？])\1+', r'\1', text)
    document = docx.Document()
    para = document.add_paragraph().add_run( 
        text)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    return document